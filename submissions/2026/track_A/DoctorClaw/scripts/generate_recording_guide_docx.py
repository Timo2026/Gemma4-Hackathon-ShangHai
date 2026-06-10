"""Generate 录屏演示指南.docx from structured content."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "录屏演示指南.docx"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_meta(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(10)


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().split("\n"):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_quote(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_title(doc, "DocClaw 录屏演示指南")
    add_meta(
        doc,
        [
            "面向产品展示、汇报录屏的分步操作与话术。技术验收步骤见 DEMO.md。",
            "演示患者：王浩然（slug: patient-zhang-san，26 岁，反复咳嗽三周）",
            "建议总时长：8–10 分钟（精简版约 4 分钟）",
            "文档版本：2026-06-07 · 对应 DocClaw Phase 5 演示能力",
        ],
    )
    doc.add_paragraph()

    add_h1(doc, "1. 录屏前准备")
    add_h2(doc, "1.1 启动环境")
    add_code_block(
        doc,
        """# 方式一：一键启动（Windows）
start.bat

# 方式二：手动四进程
cd backend && uvicorn app.main:app --reload --port 8000
cd backend && py -3.11 start_agent.py          # MCP :8001 + Agent :8090
cd frontend && npm run dev                      # :5173""",
    )
    add_para(doc, "确认浏览器可访问：http://localhost:5173")

    add_h2(doc, "1.2 配置检查")
    add_para(doc, "backend/.env 中需配置：")
    add_code_block(
        doc,
        """LLM_API_KEY=...
LLM_BASE_URL=...
LLM_MODEL=...
AGENT_API_KEY=...         # Agent 用，可与 LLM_API_KEY 相同
MONGODB_URI=mongodb://localhost:27017   # HITL 跨请求续跑（建议演示前启动）""",
    )
    add_para(doc, "HITL（人机协同确认）演示建议启动 MongoDB：")
    add_code_block(doc, "docker run -d -p 27017:27017 --name docclaw-mongo mongo:7")

    add_h2(doc, "1.3 可选预检")
    add_code_block(
        doc,
        """cd backend
py -3.11 scripts/e2e_acceptance.py --with-llm --with-agent""",
    )

    add_h2(doc, "1.4 录屏环境")
    add_table(
        doc,
        ["项目", "建议"],
        [
            ["分辨率", "1920×1080，浏览器缩放 100%"],
            ["录制范围", "仅浏览器窗口（或浏览器 + 必要说明）"],
            ["演示数据", "患者 王浩然，状态「问诊中」"],
            ["隐私", "不要录入 .env、API Key、终端报错详情"],
            ["麦克风", "录前试一句，确认音量正常"],
        ],
    )

    add_h1(doc, "2. 演示路线总览")
    add_code_block(
        doc,
        """开场定位（30s）
  → 患者队列（30s）
  → Agent 查待接诊 + 查患者（1.5min）
  → Skill 模式病历结构化（2min）★ 核心
  → Agent 创建随访 + HITL 确认（2min）★ 核心
  → 随访计划页验证（30s）
  → 通知中心（30s）
  → Agent 查待办随访（30s）
  → 技能广场收尾（30s，可选）""",
    )

    add_h1(doc, "3. 分步操作与话术")

    steps = [
        {
            "title": "【0】开场（约 30 秒）",
            "action": "操作：打开 http://localhost:5173/queue，左侧导航完整可见。",
            "script": [
                "大家好，这是 DocClaw 医疗 AI 工作台——医生的数字分身工作台。",
                "它以 Skills 为核心能力单元，配合 AI 智能体调度中枢，覆盖医生日常两类工作：实时任务（看诊、病历结构化）和计划性任务（随访、复查提醒）。",
                "接下来我用一位演示患者，走一遍完整闭环。",
            ],
        },
        {
            "title": "【1】患者队列（约 30 秒）",
            "action": "操作：停留在 /queue，指一下顶部统计卡片和患者列表，点击 王浩然 →「开始问诊」。",
            "script": [
                "首先是 患者队列。顶部可以看到今日待接诊、问诊中、已完成的统计。",
                "医生从这里进入问诊——今天演示的是 王浩然，26 岁，主诉反复咳嗽三周，状态是问诊中。",
            ],
        },
        {
            "title": "【2】Agent 模式：查待接诊 + 查患者详情（约 1.5 分钟）",
            "action": None,
            "ops": [
                "进入 /consult/patient-zhang-san",
                "右上角切换到 Agent 模式",
                "输入并发送：今天多少待接诊？",
                "等回复后，再发：王浩然的主诉和检查结果是什么？",
            ],
            "script": [
                "问诊页支持 Skill 模式 和 Agent 模式 双轨。",
                "Skill 模式是「选能力、直接调」；Agent 模式是「说人话、智能体帮你编排」。",
                "我先问 Agent：今天多少待接诊？",
                "可以看到主助理调用了患者汇总工具，返回队列统计——这是 MCP 医疗工具层 的能力。",
                "再问：王浩然的主诉和检查结果是什么？",
                "Agent 会委派 临床辅助 子智能体，自动拉取患者信息和 HIS 检查数据——医生不用切系统查。",
            ],
            "expected": "预期效果：子智能体标签显示「临床辅助」；返回主诉与检查结果。",
        },
        {
            "title": "【3】Skill 模式：门诊病历结构化（约 2 分钟）★ 核心",
            "ops": [
                "切回 Skill 模式",
                "技能下拉选择「智能病历助手」",
                "发送：请根据当前问诊整理门诊病历",
                "等 SSE 流式输出完成，指一下结构化病历卡片和字段 diff",
            ],
            "script": [
                "看诊过程中，医生切换到 Skill 模式，选择 智能病历助手。",
                "一句话：请根据当前问诊整理门诊病历。",
                "系统通过 SSE 流式返回，界面会出现 结构化病历卡片——主诉、现病史、体格检查等字段自动填充。",
                "如果有变更，还会展示 字段 diff，方便医生核对后再确认。",
                "这是 实时性任务 的典型场景：减轻文书负担，医生专注诊疗。",
            ],
            "expected": "预期效果：出现结构化病历卡片与 field_diffs 高亮。",
        },
        {
            "title": "【4】Agent 模式：创建随访 + HITL 确认（约 2 分钟）★ 核心",
            "ops": [
                "切换到 Agent 模式",
                "发送：给王浩然创建 2 周后复查的随访计划，包含复查提醒任务",
                "出现 InterruptBanner（人机协同确认条）→ 点击「确认」",
                "等待 Agent 回复创建成功",
            ],
            "script": [
                "诊后管理是 DocClaw 的另一条主线——计划性任务。",
                "我对 Agent 说：给王浩然创建 2 周后复查的随访计划，包含复查提醒任务。",
                "注意这里出现了 HITL 人机协同确认条——涉及写库、发通知等敏感操作，系统会先暂停，等医生 确认 再执行。",
                "这是医疗场景的安全设计：AI 提议，医生拍板。",
                "确认后，随访计划写入数据库，后续由后台调度器自动执行。",
            ],
            "expected": "预期效果：InterruptBanner 出现 → 点击确认 → 成功提示。",
            "note": "备注：无 MongoDB 时 HITL 仅同进程有效；正式演示建议启动 MongoDB。备用话术：「同会话内 HITL 可演示；生产环境需 MongoDB 支持跨请求续跑。」",
        },
        {
            "title": "【5】随访计划页验证（约 30 秒）",
            "action": "操作：左侧导航点击「随访计划」，找到王浩然的计划，展开任务列表。",
            "script": [
                "在 随访计划 页可以验证：计划已落库，任务包含复查提醒，状态和计划时间清晰可见。",
                "医生可以在这里手动触发，也可以交给调度器到期自动执行。",
            ],
        },
        {
            "title": "【6】通知中心（约 30 秒）",
            "action": "操作：打开 /notifications（侧边栏或顶栏入口）。",
            "script": [
                "调度器每分钟扫描到期任务，执行后会在 通知中心 推送「随访任务已执行」。",
                "这样医生不用盯后台——异常和待办主动触达，而不是被动等患者上门。",
            ],
            "note": "若通知未即时出现：录屏前可用验收脚本预创建到期任务；或话术说明「演示环境任务 2 周后到期，生产环境到期即推送」。",
            "extra_code": "cd backend\npy -3.11 scripts/e2e_acceptance.py   # 步骤 5 自动创建到期任务并验证通知",
        },
        {
            "title": "【7】Agent 查待办随访（约 30 秒）",
            "action": "操作：回到问诊页 Agent 模式，发送：有哪些待执行的随访任务？",
            "script": [
                "最后再用 Agent 问：有哪些待执行的随访任务？",
                "随访执行子智能体会汇总待办列表给医生——实时任务与计划性任务在同一块工作台闭环。",
            ],
            "expected": "预期效果：「随访执行」子智能体返回待办任务列表。",
        },
        {
            "title": "【8】收尾：技能广场（约 30 秒，可选）",
            "action": "操作：点击「技能广场」，浏览 1–2 个技能卡片。",
            "script": [
                "DocClaw 的 Skills 不只有内置能力——技能广场 支持医生获取同事分享的技能，以及 ClawHub 开放医疗技能。",
                "个人技能可启用、停用、发布，形成 可复用、可演进 的能力生态。",
                "总结一下：DocClaw 以 Skills 为核心，以 智能体调度 串联患者管理、病历文书和随访计划，让每位医生都有自己的 AI 工作台。谢谢观看。",
            ],
        },
    ]

    for step in steps:
        add_h3(doc, step["title"])
        if step.get("action"):
            add_para(doc, step["action"], bold=True)
        if step.get("ops"):
            add_para(doc, "操作：", bold=True)
            add_numbered(doc, step["ops"])
        add_para(doc, "话术：", bold=True)
        add_quote(doc, step["script"])
        if step.get("expected"):
            add_para(doc, step["expected"])
        if step.get("note"):
            add_para(doc, step["note"])
        if step.get("extra_code"):
            add_para(doc, "手动加速验收（录屏前可选）：")
            add_code_block(doc, step["extra_code"])
        doc.add_paragraph()

    add_h1(doc, "4. 精简版（4 分钟）")
    add_para(doc, "时间不够时，只录以下四段：")
    add_table(
        doc,
        ["顺序", "页面/操作", "话术要点"],
        [
            ["1", "队列 → 王浩然", "「患者管理是主线，从这里进入问诊」"],
            ["2", "Skill 模式 → 病历结构化", "「实时任务，流式输出结构化病历」"],
            ["3", "Agent 模式 → 创建随访 + HITL", "「计划性任务，AI 提议、医生确认」"],
            ["4", "随访计划页", "「计划已落库，调度器可自动执行」"],
        ],
    )
    add_para(doc, "精简版开场话术：", bold=True)
    add_quote(
        doc,
        [
            "这是 DocClaw 医疗 AI 工作台。接下来用 4 分钟演示两条主线：看诊时的 实时病历辅助，以及诊后的 随访计划与 HITL 确认。"
        ],
    )
    add_para(doc, "精简版收尾话术：", bold=True)
    add_quote(
        doc,
        ["实时看诊 + 计划随访 + 医生确认，构成 DocClaw 的核心闭环。谢谢观看。"],
    )

    add_h1(doc, "5. 录屏技巧")
    add_bullets(
        doc,
        [
            "鼠标慢、停顿 1–2 秒：等 AI 流式输出完成再解说，避免话赶画面。",
            "先干跑一遍：确认 LLM 与 Agent 响应正常，尤其步骤 3、4。",
            "关键词可后期加字幕：Skills、Agent 模式、HITL、MCP、随访调度、结构化病历。",
            "分镜建议：队列全景 → 问诊对话特写 → 病历卡片特写 → HITL 确认条特写 → 随访列表。",
            "不要录制：.env 文件、API Key、无关终端窗口。",
        ],
    )

    add_h1(doc, "6. 常见问题与备用话术")
    add_table(
        doc,
        ["现象", "处理方式", "备用话术"],
        [
            [
                "Agent API 503",
                "检查 AGENT_API_KEY，访问 :8090/health",
                "「Agent 服务需配置 API Key；Skill 模式仍可演示病历结构化。」",
            ],
            [
                "HITL 续跑失败",
                "启动 MongoDB，配置 MONGODB_URI",
                "「演示环境 HITL 需 MongoDB；逻辑上仍是医生确认后执行。」",
            ],
            [
                "Skill 无结构化输出",
                "确认技能已启用，消息含「病历」关键词",
                "「我们换一句触发语再试。」",
            ],
            [
                "通知未出现",
                "确认 Medical API 运行；或用验收脚本预创建到期任务",
                "「生产环境任务到期后自动推送通知。」",
            ],
            [
                "响应较慢",
                "录屏前预跑一遍，剪辑时可加速等待段",
                "「大模型推理需要几秒，界面会流式展示进度。」",
            ],
        ],
    )

    add_h1(doc, "7. 电梯演讲（片头/片尾）")
    add_para(doc, "中文版（约 15 秒）：", bold=True)
    add_quote(
        doc,
        [
            "DocClaw：以 Skills 为核心、以智能体为调度中枢的医生 AI 工作台——实时看诊辅助 + 计划性随访管理，AI 提议、医生确认，安全可控。"
        ],
    )
    add_para(doc, "英文版（可选）：", bold=True)
    add_quote(
        doc,
        [
            "DocClaw: A physician AI workbench powered by Skills and an agent orchestration hub — real-time clinical assistance and planned follow-up care, with human-in-the-loop safety."
        ],
    )

    add_h1(doc, "附录：页面路由速查")
    add_table(
        doc,
        ["路由", "说明"],
        [
            ["/queue", "患者队列"],
            ["/consult/patient-zhang-san", "演示患者问诊工作台"],
            ["/skills", "个人技能"],
            ["/store", "技能广场"],
            ["/followup", "随访计划"],
            ["/notifications", "通知中心"],
        ],
    )

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build()
