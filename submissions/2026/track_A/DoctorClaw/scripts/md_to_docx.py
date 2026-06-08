"""Convert project Markdown to Word (.docx) with tables and headings."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell


def set_cell_shading(cell: _Cell, fill: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_rich_text(paragraph, text: str, base_size: int = 11) -> None:
    """Parse **bold** and inline `code` in text."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, bold=True, size=base_size)
        else:
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run, size=base_size - 1)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_text(p, text, base_size=10)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "D5E8F0")
    doc.add_paragraph()


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                p.paragraph_format.left_indent = Pt(12)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if is_table_separator(line):
                i += 1
                continue
            table_rows.append(parse_table_row(line))
            i += 1
            continue

        flush_table()

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:].strip(), level=4)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            text = stripped[2:].strip()
            run = p.add_run(text)
            set_run_font(run, italic=True, size=10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x] ")
            text = stripped[6:].strip()
            p = doc.add_paragraph(style="List Bullet")
            mark = "☑ " if checked else "☐ "
            add_rich_text(p, mark + text, base_size=10)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:].strip(), base_size=10)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s", "", stripped), base_size=10)
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, stripped, base_size=11)
        i += 1

    flush_table()
    doc.save(str(docx_path))
    print(f"已生成: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md = root / "病历结构化Skill最小算力与模型配置说明.md"
    out = root / "病历结构化Skill最小算力与模型配置说明.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    convert_md_to_docx(md, out)
